# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---- Default font (Times New Roman, Unicode-friendly) ----
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = rpr.makeelement(qn('w:rFonts'), {})
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def set_font(run, size=13, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_heading(text, level=1, size=16, color=(0,0,0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_para(text, size=13, bold=False, italic=False, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_bullet(text, size=13, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level*0.6)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Consolas')
    return p

# ================= TITLE PAGE =================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BÁO CÁO BÀI TẬP LỚN')
set_font(run, size=20, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Môn học: Lập trình Web')
set_font(run, size=15, bold=True)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Đề tài: Website Flashcard Study')
set_font(run, size=14, italic=True)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Họ và tên: ........................................\nLớp: ........................................\nNhiệm vụ phụ trách: Backend, Cơ sở dữ liệu, Xác thực (Authentication)')
set_font(run, size=13)

doc.add_page_break()

# ================= PHẦN A =================
add_heading('A. LÝ DO LỰA CHỌN MIDDLEWARE TRONG MỘT TRƯỜNG HỢP CỤ THỂ', size=15)

add_heading('I. Mô tả trường hợp', level=2, size=14)

add_para('1. Chức năng cụ thể', bold=True)
add_para(
    'Trong website Flashcard Study, chức năng "Đăng nhập / Đăng ký và duy trì trạng thái '
    'đăng nhập" là một trong những chức năng quan trọng nhất, do bản thân em trực tiếp phụ trách. '
    'Cụ thể, sau khi người dùng đăng nhập thành công tại trang login.html (gửi request đến '
    'API POST /api/auth/signin), hệ thống cần phải "ghi nhớ" rằng người dùng này đã đăng nhập '
    'để các trang/API tiếp theo (ví dụ: bình luận flashcard, truy cập trang quản trị Admin) '
    'không yêu cầu đăng nhập lại.',
    indent=True
)
add_para(
    'Vấn đề đặt ra: HTTP là giao thức không trạng thái (stateless) — mỗi request gửi lên server '
    'đều độc lập, server không tự biết request này và request trước có cùng một người dùng hay không. '
    'Do đó cần một cơ chế để lưu trạng thái đăng nhập giữa nhiều request khác nhau.',
    indent=True
)

add_para('2. Lý do chọn middleware / thư viện bên thứ ba', bold=True)
add_para(
    'Để giải quyết vấn đề "duy trì trạng thái đăng nhập" nói trên, có thể tự code thủ công '
    '(ví dụ dùng localStorage ở client, hoặc tự sinh token, tự quản lý mảng lưu trạng thái '
    'trên server), nhưng các giải pháp tự viết tay sẽ gặp các nhược điểm:',
    indent=True
)
add_bullet('Tốn thời gian phát triển, dễ phát sinh lỗi bảo mật (ví dụ token bị đoán được, không có cơ chế hết hạn).')
add_bullet('Không có sẵn cơ chế sinh session ID ngẫu nhiên, an toàn.')
add_bullet('Phải tự viết lại toàn bộ cơ chế gắn/đọc cookie ở mỗi route, code lặp lại nhiều.')
add_para(
    'Vì vậy, nhóm quyết định sử dụng middleware có sẵn là thư viện express-session — một thư viện '
    'middleware phổ biến, được cộng đồng Node.js/Express tin dùng, đã được kiểm thử kỹ và tích hợp '
    'rất dễ dàng vào Express thông qua app.use().',
    indent=True
)

add_para('3. Middleware được chọn', bold=True)

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
rows_data = [
    ('Tên middleware', 'express-session (thư viện bên thứ 3 của Express)'),
    ('Vai trò', 'Quản lý phiên làm việc (session) của người dùng trên server'),
    ('Nhiệm vụ',
     'Tạo session ID ngẫu nhiên cho mỗi client; lưu dữ liệu người dùng (req.session.user) '
     'trên server; tự động gắn cookie chứa session ID vào response và đọc lại cookie đó '
     'ở các request tiếp theo để khôi phục dữ liệu session.'),
    ('Vị trí sử dụng trong code', 'server.js — app.use(session({ secret, resave, saveUninitialized }))'),
    ('Lý do chọn',
     '(1) Tích hợp sẵn với Express, chỉ cần 1 dòng app.use(); '
     '(2) Tự động sinh session ID an toàn, không cần tự viết; '
     '(3) Cho phép lưu trực tiếp object user (req.session.user) mà không cần parse/encode token; '
     '(4) Phù hợp với quy mô bài tập lớn (không cần JWT phức tạp); '
     '(5) Dễ kết hợp với middleware tự viết isAuthenticated/isAdmin để phân quyền.'),
]
for i, (k, v) in enumerate(rows_data):
    table.cell(i,0).text = k
    table.cell(i,1).text = v
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, size=12)

doc.add_paragraph()
add_code(
    'app.use(session({\n'
    "  secret: 'autosecretkey',\n"
    '  resave: false,\n'
    '  saveUninitialized: true\n'
    '}));'
)

doc.add_page_break()

# ================= PHẦN II =================
add_heading('II. Giải thích đường đi từ Request đến Response', size=15)

add_para(
    'Lấy ví dụ cụ thể chức năng "Đăng nhập" để minh họa đường đi của một request, '
    'trong đó middleware express-session đóng vai trò trung gian giữa Controller và Client.'
)

steps = [
    ('Bước 1 — Client gửi request',
     'Người dùng nhập username/password tại trang login.html, JavaScript (public/js/login.js) '
     'gửi: POST /api/auth/signin với body {username, password} dạng JSON.'),
    ('Bước 2 — Middleware session xử lý trước (vào)',
     'Trước khi request đến Controller, middleware express-session (đăng ký toàn cục trong server.js) '
     'chạy đầu tiên: đọc cookie connect.sid trong request (nếu có) để xác định session cũ, '
     'hoặc tạo session mới và gắn object req.session vào request.'),
    ('Bước 3 — Express Router định tuyến',
     "app.use('/api/auth', authRoutes) khớp tiền tố URL, chuyển tiếp sang authRoutes.js. "
     "Tại đây router.post('/signin', signin) khớp đúng phương thức POST và path /signin, "
     "gọi hàm signin() trong authController.js."),
    ('Bước 4 — Controller xử lý nghiệp vụ + truy vấn Database',
     "Controller thực hiện: (a) SELECT user theo username từ bảng users trong MySQL; "
     "(b) dùng bcrypt.compare() so sánh mật khẩu nhập vào với mật khẩu hash lưu trong DB; "
     "(c) nếu đúng, gán req.session.user = {id, username, role}."),
    ('Bước 5 — Middleware session xử lý sau (ra)',
     'Khi Controller gọi res.json(), middleware session lưu lại dữ liệu req.session.user vào bộ nhớ '
     'session trên server, đồng thời tự động đính kèm header Set-Cookie chứa session ID vào response.'),
    ('Bước 6 — Response trả về Client',
     'Client nhận HTTP 200 với JSON {ok: true, user: {...}} kèm cookie session. Trình duyệt tự lưu '
     'cookie này; mọi request sau đó (ví dụ gửi bình luận, vào trang Admin) sẽ tự động gửi kèm cookie, '
     'giúp server nhận diện đúng người dùng mà không cần đăng nhập lại.'),
]

for title_step, desc in steps:
    add_para(title_step, bold=True)
    add_para(desc, indent=True)

add_para('Sơ đồ tổng quát:', bold=True)
add_code(
    'Browser\n'
    '   |  POST /api/auth/signin {username, password}\n'
    '   v\n'
    'server.js --> [Middleware: express-session] --> đọc/tạo session\n'
    '   |\n'
    '   v\n'
    'authRoutes.js (router.post("/signin", signin))\n'
    '   |\n'
    '   v\n'
    'authController.js: signin()\n'
    '   |---> MySQL: SELECT * FROM users WHERE username = ?\n'
    '   |---> bcrypt.compare(password, hash)\n'
    '   |---> req.session.user = {id, username, role}\n'
    '   v\n'
    '[Middleware: express-session] lưu session + gắn Set-Cookie\n'
    '   v\n'
    'Response: {ok:true, user:{...}}  -->  Browser (lưu cookie)'
)

doc.add_page_break()

# ================= PHẦN III =================
add_heading('III. Giải thích vận hành website, tập trung vào phần việc được phân công', size=15)

add_para('1. Tổng quan phần việc được phân công', bold=True)
add_para(
    'Trong dự án Flashcard Study, em được phân công phụ trách phần BACKEND, cụ thể gồm: '
    'thiết kế và quản lý Cơ sở dữ liệu (MySQL), xây dựng hệ thống Xác thực người dùng '
    '(đăng ký/đăng nhập/đăng xuất), middleware phân quyền Admin, và tổ chức cấu trúc thư mục '
    'dự án theo mô hình MVC (Model - View - Controller) để các thành viên khác trong nhóm '
    'dễ dàng phát triển phần Frontend song song.',
    indent=True
)

add_para('2. Vận hành phần Cơ sở dữ liệu (MySQL)', bold=True)
add_para(
    'File src/config/database.js tạo một Connection Pool (tối đa 10 kết nối đồng thời) đến '
    'MySQL, giúp xử lý nhiều request truy vấn cùng lúc mà không cần mở/đóng kết nối liên tục. '
    'Khi server khởi động, hàm initDatabase() tự động kiểm tra và tạo các bảng users, comments, '
    'contents nếu chưa tồn tại, đồng thời tạo sẵn 1 tài khoản admin mặc định và một số dữ liệu mẫu.',
    indent=True
)
add_bullet('Bảng users: lưu tài khoản (username, email, password đã mã hóa bcrypt, role admin/user).')
add_bullet('Bảng contents: lưu nội dung flashcard (title, body, type).')
add_bullet('Bảng comments: lưu bình luận và đánh giá của người dùng theo từng flashcard (contentId).')
add_bullet('Bảng page_views, view_logs: thống kê lượt truy cập trang.')

add_para('3. Vận hành hệ thống Xác thực (Authentication)', bold=True)
add_para(
    'File src/controllers/authController.js xử lý 4 chức năng chính:',
    indent=True
)
add_bullet('signup: kiểm tra trùng username/email, mã hóa mật khẩu bằng bcrypt trước khi lưu vào DB.')
add_bullet('signin: so sánh mật khẩu (bcrypt.compare), nếu đúng thì lưu thông tin vào session.')
add_bullet('me: kiểm tra session hiện tại để biết người dùng đã đăng nhập hay chưa.')
add_bullet('signout: hủy session (req.session.destroy()), đăng xuất khỏi hệ thống.')

add_para('4. Vận hành Middleware phân quyền', bold=True)
add_para(
    'File src/middleware/auth.js cung cấp 2 hàm chặn (guard) request trước khi vào Controller:',
    indent=True
)
add_bullet('isAuthenticated: nếu chưa có req.session.user thì chuyển hướng (redirect) về trang đăng nhập.')
add_bullet('isAdmin: nếu req.session.user.role khác "admin" thì trả về lỗi 403 Forbidden, ngăn người '
           'dùng thường truy cập trang quản trị (/admin).')
add_para(
    'Hai middleware này được gắn vào route /admin (src/routes/admin.js) để bảo vệ toàn bộ chức năng '
    'quản trị nội dung flashcard mà chỉ admin được quyền chỉnh sửa.',
    indent=True
)

add_para('5. Kết quả vận hành tổng thể', bold=True)
add_para(
    'Nhờ middleware express-session kết hợp với middleware tự viết isAuthenticated/isAdmin, website '
    'đạt được luồng hoạt động an toàn và mạch lạc: người dùng đăng nhập một lần, trạng thái được duy trì '
    'xuyên suốt qua các trang; người dùng thường chỉ xem được flashcard và bình luận; chỉ admin mới '
    'truy cập được trang quản trị để thêm/sửa/xoá nội dung flashcard trong cơ sở dữ liệu.',
    indent=True
)

doc.save('BAOCAO_MIDDLEWARE.docx')
print('Saved BAOCAO_MIDDLEWARE.docx')
